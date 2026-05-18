"""Generate LeanChem_Connect_User_Guide.docx from USER_BLUEPRINT.md."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "USER_BLUEPRINT.md"
OUT = ROOT / "LeanChem_Connect_User_Guide.docx"

EMOJI_RE = re.compile(
    "[\U0001F300-\U0001FAFF\U00002600-\U000027BF\U0001F000-\U0001F2FF]+",
    flags=re.UNICODE,
)

def clean(text: str) -> str:
    text = EMOJI_RE.sub("", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def add_inline(paragraph, text: str):
    """Render **bold**, `code`, [link](url) as runs in a paragraph."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            r = paragraph.add_run(label)
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def main():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LeanChem Connect")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Complete User Guide")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Integrated Deal and Product System (IDPS)\nVersion 1.0")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()

    in_code = False
    code_buffer = []
    in_list_kind = None  # 'ul' or 'ol' or None

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        p = doc.add_paragraph()
        r = p.add_run("\n".join(code_buffer))
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        code_buffer = []

    for raw in md:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_buffer.append(raw)
            continue

        if not line.strip():
            in_list_kind = None
            continue

        if line.strip() == "---":
            in_list_kind = None
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = clean(m.group(2))
            if not text:
                continue
            heading = doc.add_heading(text, level=min(level, 4))
            for run in heading.runs:
                if level == 1:
                    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)
                elif level == 2:
                    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
                else:
                    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            in_list_kind = None
            continue

        # Bulleted list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            text = clean(m.group(2))
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent)
            add_inline(p, text)
            in_list_kind = "ul"
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            text = clean(m.group(2))
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent)
            add_inline(p, text)
            in_list_kind = "ol"
            continue

        # Plain paragraph
        text = clean(line)
        if not text:
            continue
        p = doc.add_paragraph()
        add_inline(p, text)
        in_list_kind = None

    flush_code()

    doc.save(OUT)
    print(f"Wrote {OUT}")

if __name__ == "__main__":
    main()
